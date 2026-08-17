"""Does SURFACING actually produce a recognisable pick-list from the real corpus?
Union-find families + shared facts + sessions. No model. The user ticks boxes."""
import re, time, hashlib
from pathlib import Path
from collections import defaultdict, Counter

DL = Path.home() / "Downloads"
import pypdfium2 as pdfium
from docx import Document

files = [p for p in DL.iterdir() if p.is_file() and p.name != ".DS_Store"]
DOCS = [p for p in files if p.suffix.lower() in (".pdf", ".docx")][:400]
print(f"corpus {len(files)} files · reading {len(DOCS)} docs\n", flush=True)

# ---------- union-find ----------
par = {}
def find(x):
    par.setdefault(x, x)
    while par[x] != x: par[x] = par[par[x]]; x = par[x]
    return x
def union(a, b):
    ra, rb = find(a), find(b)
    if ra != rb: par[ra] = rb

# 1. version/duplicate families by normalised stem
VER = re.compile(r'\s*(\(\d+\)|copy( \d+)?|v\d+|final\d*|updated?|draft\d*|\d+)\s*$', re.I)
def stem_key(n):
    s, prev = Path(n).stem, None
    while s != prev: prev = s; s = VER.sub("", s).strip()
    return s.lower()
bystem = defaultdict(list)
for p in files: bystem[stem_key(p.name)].append(p)
for k, v in bystem.items():
    if len(v) > 1:
        for q in v[1:]: union(str(v[0]), str(q))

# 2. shared content facts
ENTP = {
    "code":  re.compile(r'\b([A-Z]{2,5})\s?-?\s?(\d{3,4})\b'),
    "email": re.compile(r'\b[\w.+-]+@([\w-]+\.[\w.-]+)\b'),
    "uniid": re.compile(r'\b([a-z]{2,4}\d{4})\b'),
}
JUNK = {"gmail.com","outlook.com","hotmail.com","yahoo.com","icloud.com","w3.org",
        "adobe.com","openxmlformats.org","googlemail.com"}
def read(p):
    try:
        if p.suffix.lower()==".pdf":
            d=pdfium.PdfDocument(str(p)); t=[d[i].get_textpage().get_text_range() for i in range(min(len(d),6))]
            d.close(); return "".join(t)
        doc=Document(str(p)); return " ".join(x.text for x in doc.paragraphs[:80])
    except Exception: return None

t0=time.time(); ent=defaultdict(set)
for p in DOCS:
    t=read(p)
    if not t: continue
    hay=f"{p.name}\n{t[:6000]}"
    for m in ENTP["code"].finditer(hay):  ent[("code",f"{m.group(1)}{m.group(2)}")].add(str(p))
    for m in ENTP["email"].finditer(hay):
        d=m.group(1).lower()
        if d not in JUNK: ent[("org",d)].add(str(p))
    for m in ENTP["uniid"].finditer(hay): ent[("id",m.group(1))].add(str(p))
print(f"read in {time.time()-t0:.0f}s", flush=True)

N=len(DOCS)
FACT={e:s for e,s in ent.items() if 3<=len(s)<=int(0.30*N)}   # IDF band + hub cap
for e,s in FACT.items():
    s=sorted(s)
    for q in s[1:]: union(s[0],q)

# 3. download sessions (purpose signal, INFERRED)
mt={str(p):p.stat().st_mtime for p in files}
srt=sorted(mt,key=mt.get)
sess=[]; cur=[srt[0]]
for a,b in zip(srt,srt[1:]):
    if mt[b]-mt[a]<=600: cur.append(b)
    else:
        if len(cur)>=4: sess.append(cur)
        cur=[b]
if len(cur)>=4: sess.append(cur)

# ---------- assemble candidate groups ----------
comp=defaultdict(list)
for p in files: comp[find(str(p))].append(p)
groups=[(v,"facts/versions") for v in comp.values() if len(v)>=4]
groups+= [([Path(x) for x in s],"one download session") for s in sess if len(s)>=6]

def label(members):
    toks=Counter()
    for p in members:
        for w in re.split(r'[^A-Za-z0-9]+', Path(p).stem):
            if len(w)>2 and not w.isdigit(): toks[w.lower()]+=1
    top=[w for w,c in toks.most_common(4) if c>=max(2,len(members)*0.25)]
    return " · ".join(top) if top else "(mixed)"

groups.sort(key=lambda g:-len(g[0]))
print(f"\n{'='*74}\nSURFACED: {len(groups)} candidate groups — user ticks boxes\n{'='*74}\n")
seen=set(); shown=0
for members,why in groups:
    key=frozenset(str(m) for m in members)
    if key in seen: continue
    seen.add(key); shown+=1
    if shown>16: break
    exts=Counter(p.suffix.lower().lstrip('.') for p in members)
    print(f"  ☐ {label(members):<44} {len(members):>4} files   [{why}]")
    print(f"      {', '.join(f'{k}:{v}' for k,v in exts.most_common(4))}")
    for p in members[:3]: print(f"      · {p.name[:66]}")
    print(flush=True)
