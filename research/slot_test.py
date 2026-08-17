"""Can we fill template SLOTS from real file content? Validated extraction, not naive regex."""
import re, time
from pathlib import Path
from collections import Counter

DL = Path.home() / "Downloads"
import pypdfium2 as pdfium
from docx import Document

# ---------- gazetteers: closed-ish sets make validation possible ----------
SCHOOLS = {
    "columbia": "Columbia", "barnard": "Barnard", "yale": "Yale", "stanford": "Stanford",
    "harvard": "Harvard", "northwestern": "Northwestern", "duke": "Duke", "dartmouth": "Dartmouth",
    "uchicago": "UChicago", "university of chicago": "UChicago", "cornell": "Cornell",
    "johns hopkins": "Johns Hopkins", "georgetown": "Georgetown", "wash u": "WashU",
    "washington university": "WashU", "hku": "HKU", "university of hong kong": "HKU",
    "unc": "UNC", "princeton": "Princeton", "mit": "MIT", "nyu": "NYU", "upenn": "UPenn",
}
WORKTYPE = {
    "syllabus": "Syllabus", "lecture": "Lecture", "problem set": "Problem Set", "pset": "Problem Set",
    "homework": "Homework", "assignment": "Assignment", "exam": "Exam", "midterm": "Exam",
    "final exam": "Exam", "quiz": "Quiz", "practice": "Practice", "solution": "Solutions",
    "essay": "Essay", "lab report": "Lab Report", "notes": "Notes", "transcript": "Transcript",
    "application": "Application", "resume": "Resume", "cover letter": "Cover Letter",
    "abstract": "Abstract", "manuscript": "Manuscript", "invoice": "Invoice", "receipt": "Receipt",
}
# subject code: LETTERS+DIGITS *validated by academic context nearby*
SUBJ = re.compile(r'\b([A-Z]{2,5})\s?-?\s?(\d{3,4})(?:\.\d{1,4})?\b')
ACADEMIC_CTX = re.compile(
    r'\b(course|section|syllabus|lecture|semester|instructor|professor|credits|prerequisite|'
    r'department|spring|fall|summer|winter|registrar|enroll|class number|call number)\b', re.I)
TERM = re.compile(r'\b(Spring|Fall|Autumn|Summer|Winter)\s*(20\d{2})\b', re.I)
YEAR = re.compile(r'\b(20[0-2]\d)\b')


def read(p):
    try:
        if p.suffix.lower() == ".pdf":
            d = pdfium.PdfDocument(str(p)); t = []
            for i in range(len(d)):
                t.append(d[i].get_textpage().get_text_range())
            d.close(); return "".join(t)
        if p.suffix.lower() == ".docx":
            doc = Document(str(p))
            body = " ".join(x.text for x in doc.paragraphs)
            for tb in doc.tables:
                for r in tb.rows: body += " " + " ".join(c.text for c in r.cells)
            return body
    except Exception:
        return None


def slots(name, text):
    """Fill Academic template slots. Every slot must be VALIDATED, never merely matched."""
    s, ev = {}, {}
    hay = f"{name}\n{text}"
    low = hay.lower()

    for k, v in SCHOOLS.items():                       # gazetteer -> closed set
        if k in low:
            s["school"] = v; ev["school"] = f"gazetteer:{k}"; break

    m = TERM.search(hay)                               # explicit term beats bare year
    if m:
        s["term"] = f"{m.group(2)}-{m.group(1).title()}"; ev["term"] = f"text:{m.group(0)}"
    else:
        yrs = YEAR.findall(hay)
        if yrs:
            s["term"] = Counter(yrs).most_common(1)[0][0]; ev["term"] = "year-mode"

    # subject code: require academic context within +/-300 chars -> kills VHX7000
    for m in SUBJ.finditer(hay):
        a, b = max(0, m.start() - 300), min(len(hay), m.end() + 300)
        if ACADEMIC_CTX.search(hay[a:b]):
            s["subject"] = f"{m.group(1)}{m.group(2)}"; ev["subject"] = f"code+ctx:{m.group(0)}"
            break

    for k, v in WORKTYPE.items():                      # prefer filename evidence
        if k in name.lower():
            s["work_type"] = v; ev["work_type"] = f"filename:{k}"; break
    else:
        for k, v in WORKTYPE.items():
            if k in low[:4000]:
                s["work_type"] = v; ev["work_type"] = f"text:{k}"; break
    return s, ev


files = [p for p in DL.iterdir()
         if p.is_file() and p.suffix.lower() in (".pdf", ".docx")][:300]
print(f"testing {len(files)} real files\n", flush=True)

t0 = time.time()
filled = Counter(); rows = []; n_ok = 0
for p in files:
    t = read(p)
    if t is None: continue
    n_ok += 1
    s, ev = slots(p.name, t)
    for k in s: filled[k] += 1
    if len(s) >= 3: rows.append((p.name, s, ev))
dt = time.time() - t0

print(f"read+extracted {n_ok} files in {dt:.1f}s ({1000*dt/max(n_ok,1):.0f} ms/file)\n")
print("=== SLOT FILL RATE (Academic template) ===")
for k in ("school", "term", "subject", "work_type"):
    print(f"  {k:<10} {filled[k]:4}/{n_ok}  ({100*filled[k]/max(n_ok,1):.0f}%)")
print(f"\n  files with >=3 of 4 slots: {len(rows)} ({100*len(rows)/max(n_ok,1):.0f}%)\n", flush=True)

print("=== REAL PLACEMENTS (school/term/subject/work_type) ===\n")
for name, s, ev in rows[:18]:
    path = "/".join(s.get(k, "·") for k in ("school", "term", "subject", "work_type"))
    print(f"{path}")
    print(f"    ← {name[:70]}")
    print(f"      {ev}\n", flush=True)
