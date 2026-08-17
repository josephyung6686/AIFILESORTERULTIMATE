"""Does splitting `school` into authored_by / addressed_to actually resolve the failures?
Hypothesis: filename+prompt-language = addressed_to ; identity markers = authored_by."""
import re
from pathlib import Path
from collections import defaultdict

DL = Path.home() / "Downloads"
import pypdfium2 as pdfium
from docx import Document

SCHOOLS = {
    "columbia": "Columbia", "barnard": "Barnard", "yale": "Yale", "stanford": "Stanford",
    "harvard": "Harvard", "northwestern": "Northwestern", "duke": "Duke", "dartmouth": "Dartmouth",
    "uchicago": "UChicago", "u chicago": "UChicago", "university of chicago": "UChicago",
    "cornell": "Cornell", "johns hopkins": "Johns Hopkins", "georgetown": "Georgetown",
    "wash u": "WashU", "washu": "WashU", "washington university": "WashU", "hku": "HKU",
    "unc": "UNC", "princeton": "Princeton", "mit": "MIT", "nyu": "NYU", "upenn": "UPenn",
    "morehead": "UNC", "brown": "Brown", "rice": "Rice", "vanderbilt": "Vanderbilt",
}
def wb(k): return re.compile(r'(?<![a-z0-9])' + re.escape(k) + r'(?![a-z0-9])', re.I)
SRE = {k: (wb(k), v) for k, v in SCHOOLS.items()}

# identity markers = who WROTE it
IDENT = re.compile(r'([\w.+-]+)@([\w-]+\.edu)\b|\b([a-z]{2,4}\d{4})\b', re.I)
# second-person prompt language = who it is ADDRESSED to
PROMPT = re.compile(r'\b(why (do you|are you|us|our)|what (do you|attracts you)|'
                    r'how will you|our (community|campus|program|university)|'
                    r'attend (our|this)|supplement|supplemental|applicant)\b', re.I)


def read(p):
    try:
        if p.suffix.lower() == ".pdf":
            d = pdfium.PdfDocument(str(p)); t = [d[i].get_textpage().get_text_range() for i in range(len(d))]
            d.close(); return "".join(t)
        doc = Document(str(p))
        return " ".join(x.text for x in doc.paragraphs)
    except Exception: return None


def schools_in(s):
    out = defaultdict(int)
    for k, (rx, v) in SRE.items():
        n = len(rx.findall(s))
        if n: out[v] += n
    return out


files = [p for p in DL.iterdir() if p.is_file() and p.suffix.lower() in (".pdf", ".docx")]
rows = []
for p in files:
    fn_sch = schools_in(p.name)
    if not fn_sch: continue                      # need a filename school to test the split
    t = read(p)
    if not t: continue
    body_sch = schools_in(t[:9000])
    if not body_sch: continue

    addressed = max(fn_sch, key=fn_sch.get)      # filename = target

    # authored_by: school appearing beside an identity marker
    authored = None
    for m in IDENT.finditer(t[:9000]):
        span = t[max(0, m.start()-120): m.end()+120]
        s = schools_in(span)
        if s: authored = max(s, key=s.get); break
    if not authored:                             # fallback: most frequent body school ≠ target
        others = {k: v for k, v in body_sch.items() if k != addressed}
        if others: authored = max(others, key=others.get)

    is_app = bool(PROMPT.search(t[:6000])) or "supplement" in p.name.lower()
    if authored and authored != addressed:
        rows.append((p.name, addressed, authored, is_app))

print(f"files where FILENAME school != IDENTITY school: {len(rows)}\n")
print(f"{'file':<50} {'addressed_to':<14} {'authored_by':<14} app?")
print("-" * 88)
for n, a, w, isapp in rows[:26]:
    print(f"{n[:48]:<50} {a:<14} {w:<14} {'yes' if isapp else ''}")

apps = sum(1 for r in rows if r[3])
print(f"\n{apps}/{len(rows)} carry application-prompt language → the disagreement IS the role split")
