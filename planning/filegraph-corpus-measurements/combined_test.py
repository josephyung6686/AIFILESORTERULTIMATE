"""The number that matters: what does rules + propagation actually cover, and is it right?
Propagate into the files where the disciplined extractor ABSTAINED, then show the output for
human judgement. This is the pre-LLM floor of the system."""
import re, time, math
from pathlib import Path
from collections import defaultdict, Counter

DL = Path.home() / "Downloads"
import pypdfium2 as pdfium
from docx import Document

SCHOOLS = {
    "columbia": "Columbia", "barnard": "Barnard", "yale": "Yale", "stanford": "Stanford",
    "harvard": "Harvard", "northwestern": "Northwestern", "duke": "Duke", "dartmouth": "Dartmouth",
    "uchicago": "UChicago", "u chicago": "UChicago", "cornell": "Cornell",
    "johns hopkins": "Johns Hopkins", "georgetown": "Georgetown", "wash u": "WashU",
    "washu": "WashU", "hku": "HKU", "unc": "UNC", "princeton": "Princeton", "mit": "MIT",
    "nyu": "NYU", "upenn": "UPenn", "brown": "Brown", "rice": "Rice",
}
def wb(k): return re.compile(r'(?<![a-z0-9])' + re.escape(k) + r'(?![a-z0-9])', re.I)
SRE = {k: (wb(k), v) for k, v in SCHOOLS.items()}
ENT = {
    "email": re.compile(r'\b[\w.+-]+@([\w-]+\.[\w.-]+)\b'),
    "code":  re.compile(r'\b([A-Z]{2,5})\s?-?\s?(\d{3,4})\b'),
    "uniid": re.compile(r'\b([a-z]{2,4}\d{4})\b'),
}
JUNK = {"gmail.com", "outlook.com", "hotmail.com", "yahoo.com", "icloud.com",
        "w3.org", "adobe.com", "openxmlformats.org", "googlemail.com"}


def read(p):
    try:
        if p.suffix.lower() == ".pdf":
            d = pdfium.PdfDocument(str(p)); t = [d[i].get_textpage().get_text_range() for i in range(len(d))]
            d.close(); return "".join(t)
        doc = Document(str(p))
        b = " ".join(x.text for x in doc.paragraphs)
        for tb in doc.tables:
            for r in tb.rows: b += " " + " ".join(c.text for c in r.cells)
        return b
    except Exception: return None


def rules_school(name, text):
    sc = defaultdict(float)
    for zt, w in ((name, 10.0), (text[:1200], 3.0), (text[1200:], 0.4)):
        for k, (rx, v) in SRE.items():
            n = len(rx.findall(zt))
            if n: sc[v] += w * (1 + 0.15 * min(n, 6))
    if not sc: return None
    r = sorted(sc.items(), key=lambda kv: -kv[1])
    top = r[0][1]; sec = r[1][1] if len(r) > 1 else 0.0
    return r[0][0] if (top >= 3.0 and top - sec >= 1.0) else None


def ents(name, text):
    out = set(); hay = f"{name}\n{text[:8000]}"
    for m in ENT["email"].finditer(hay):
        d = m.group(1).lower()
        if d not in JUNK: out.add(("email", d))
    for m in ENT["code"].finditer(hay):  out.add(("code", f"{m.group(1)}{m.group(2)}"))
    for m in ENT["uniid"].finditer(hay): out.add(("uniid", m.group(1)))
    return out


files = [p for p in DL.iterdir() if p.is_file() and p.suffix.lower() in (".pdf", ".docx")][:400]
print(f"reading {len(files)}…", flush=True)
t0 = time.time()
R, E, MT, NM = {}, {}, {}, {}
for i, p in enumerate(files):
    t = read(p)
    if t is None: continue
    NM[i] = p.name; MT[i] = p.stat().st_mtime
    R[i] = rules_school(p.name, t); E[i] = ents(p.name, t)
idx = list(NM)
n_rules = sum(1 for i in idx if R[i])
print(f"read in {time.time()-t0:.0f}s\n")
print(f"RULES alone      : {n_rules}/{len(idx)} = {100*n_rules/len(idx):.0f}% coverage\n", flush=True)

df = Counter()
for i in idx: df.update(E[i])
N = len(idx)
KEEP = {e for e, c in df.items() if 2 <= c <= max(3, int(0.25 * N))}
inv = defaultdict(list)
for i in idx:
    for e in E[i] & KEEP: inv[e].append(i)
W = defaultdict(float)
for e, fl in inv.items():
    w = math.log(N / len(fl))
    for a in range(len(fl)):
        for b in range(a+1, len(fl)): W[(fl[a], fl[b])] += w
srt = sorted(idx, key=lambda i: MT[i])
for a in range(len(srt)-1):
    for b in range(a+1, min(a+6, len(srt))):
        if abs(MT[srt[a]] - MT[srt[b]]) <= 600:
            W[(min(srt[a],srt[b]), max(srt[a],srt[b]))] += 0.35
adj = defaultdict(list)
for (a, b), w in W.items(): adj[a].append((b, w)); adj[b].append((a, w))

# propagate into the abstainers
gained = []
for i in idx:
    if R[i]: continue
    vote = defaultdict(float)
    for j, w in adj[i]:
        if R[j]: vote[R[j]] += w
    if not vote: continue
    r = sorted(vote.items(), key=lambda kv: -kv[1])
    pred, top = r[0]; sec = r[1][1] if len(r) > 1 else 0.0
    if top >= 1.0 and (top - sec) >= 0.5 * top:
        gained.append((NM[i], pred, round(top, 1)))

tot = n_rules + len(gained)
print(f"+ PROPAGATION    : +{len(gained)} files")
print(f"COMBINED         : {tot}/{len(idx)} = {100*tot/len(idx):.0f}% coverage")
print(f"                   (rules {100*n_rules/len(idx):.0f}% → combined {100*tot/len(idx):.0f}%)\n")
print("=== what propagation added — judge these by eye ===\n")
for nm, pred, w in gained[:30]:
    print(f"  {pred:<14} ← {nm[:62]}   (w={w})")
