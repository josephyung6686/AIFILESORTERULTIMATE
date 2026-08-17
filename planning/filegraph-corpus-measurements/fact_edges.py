"""Does content extraction produce SHARED FACTS that link real files? The core claim, tested."""
import re, time, json
from pathlib import Path
from collections import defaultdict, Counter

DL = Path.home() / "Downloads"
import pypdfium2 as pdfium
from docx import Document

# --- entity patterns: things two files can SHARE as an identity, not a similarity ---
PAT = {
    "course":  re.compile(r'\b([A-Z]{2,5})\s?-?\s?(\d{4})(?:\.\d{1,4})?\b'),
    "doi":     re.compile(r'\b10\.\d{4,9}/[-._;()/:A-Za-z0-9]+'),
    "email":   re.compile(r'\b[\w.+-]+@([\w-]+\.[\w.-]+)\b'),
    "uni_id":  re.compile(r'\b([a-z]{2,4}\d{4})\b'),
    "url_host":re.compile(r'https?://([\w.-]+)'),
    "money_yr":re.compile(r'\b(20[12]\d)\b'),
}
STOP_HOST = {"www.w3.org", "schemas.openxmlformats.org", "purl.org", "ns.adobe.com"}


def text_of(p):
    ext = p.suffix.lower()
    try:
        if ext == ".pdf":
            d = pdfium.PdfDocument(str(p)); meta = {}
            for k in ("Author", "Title", "Creator"):
                v = d.get_metadata_value(k)
                if v: meta[k] = v.strip()
            t = ""
            for i in range(min(2, len(d))):
                t += d[i].get_textpage().get_text_range()
            d.close()
            return t[:6000], meta
        if ext == ".docx":
            doc = Document(str(p)); cp = doc.core_properties
            meta = {k: str(getattr(cp, k)).strip() for k in ("author", "title", "company")
                    if getattr(cp, k, None)}
            t = " ".join(x.text for x in doc.paragraphs[:60])
            return t[:6000], meta
    except Exception:
        return None, None
    return None, None


files = [p for p in DL.iterdir() if p.is_file() and p.suffix.lower() in (".pdf", ".docx")]
SAMPLE = files[:400]
print(f"sampling {len(SAMPLE)} of {len(files)} PDF/DOCX\n", flush=True)

ents = defaultdict(set)        # entity -> {file idx}
per_file = defaultdict(set)    # file idx -> {entity}
names, ok, t0 = {}, 0, time.time()

for i, p in enumerate(SAMPLE):
    t, meta = text_of(p)
    if t is None: continue
    ok += 1; names[i] = p.name
    found = set()
    for m in PAT["course"].finditer(t):
        found.add(("course", f"{m.group(1)}{m.group(2)}"))
    for m in PAT["doi"].finditer(t):
        found.add(("doi", m.group(0)[:40]))
    for m in PAT["email"].finditer(t):
        h = m.group(1).lower()
        if h not in STOP_HOST: found.add(("emaildom", h))
    for m in PAT["uni_id"].finditer(t):
        found.add(("uniid", m.group(1)))
    for m in PAT["url_host"].finditer(t):
        h = m.group(1).lower()
        if h not in STOP_HOST: found.add(("host", h))
    for k, v in (meta or {}).items():
        v = v.strip()
        if v and v.lower() not in ("python-docx", "word document", "untitled", "microsoft word"):
            found.add((f"meta_{k.lower()}", v[:60]))
    for e in found:
        ents[e].add(i); per_file[i].add(e)

dt = time.time() - t0
print(f"extracted {ok}/{len(SAMPLE)} files in {dt:.1f}s ({1000*dt/max(ok,1):.0f} ms/file)\n", flush=True)

# --- how many entities are SHARED (link >=2 files) and how many edges do they make? ---
shared = {e: fs for e, fs in ents.items() if 2 <= len(fs) <= 60}
edges = set()
for e, fs in shared.items():
    fl = sorted(fs)
    for a in range(len(fl)):
        for b in range(a + 1, len(fl)):
            edges.add((fl[a], fl[b]))

linked = {i for ed in edges for i in ed}
print(f"distinct entities extracted : {len(ents):,}")
print(f"SHARED entities (link 2-60 files): {len(shared):,}")
print(f"FACT EDGES produced         : {len(edges):,}")
print(f"files with >=1 fact edge    : {len(linked)} / {ok}  ({100*len(linked)/max(ok,1):.0f}%)\n", flush=True)

by_kind = Counter(e[0] for e in shared)
print("shared entities by kind:", dict(by_kind.most_common()), "\n", flush=True)

print("=== strongest shared facts (what actually links your files) ===\n")
for e, fs in sorted(shared.items(), key=lambda kv: -len(kv[1]))[:16]:
    if len(fs) > 25: continue
    print(f"[{e[0]}] {e[1]!r}  links {len(fs)} files:")
    for i in sorted(fs)[:4]:
        print(f"      {names[i][:74]}")
    print(flush=True)

json.dump({"sampled": len(SAMPLE), "ok": ok, "ms_per_file": round(1000*dt/max(ok,1)),
           "entities": len(ents), "shared": len(shared), "edges": len(edges),
           "linked_pct": round(100*len(linked)/max(ok,1))},
          open("fact_edges_result.json", "w"), indent=2)
