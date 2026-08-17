"""How expensive is REAL content extraction on the real corpus? Facts, not similarity."""
import time, json, re, sys
from pathlib import Path
from collections import Counter

DL = Path.home() / "Downloads"
files = [p for p in DL.iterdir() if p.is_file() and p.name != ".DS_Store"]

# ---------- PDF ----------
import pypdfium2 as pdfium

def pdf_facts(p):
    f = {}
    doc = pdfium.PdfDocument(str(p))
    f["pages"] = len(doc)
    meta = {}
    for k in ("Title", "Author", "Subject", "Creator", "Producer", "CreationDate"):
        try:
            v = doc.get_metadata_value(k)
            if v: meta[k] = v[:120]
        except Exception: pass
    f["meta"] = meta
    tp = doc[0].get_textpage()
    f["text"] = tp.get_text_range()[:1500]
    doc.close()
    return f

# ---------- DOCX ----------
from docx import Document

def docx_facts(p):
    d = Document(str(p))
    cp = d.core_properties
    meta = {k: str(getattr(cp, k))[:120] for k in ("title", "author", "company", "subject", "last_modified_by")
            if getattr(cp, k, None)}
    heads = [x.text.strip() for x in d.paragraphs if x.style.name.startswith("Heading") and x.text.strip()][:8]
    body = " ".join(x.text for x in d.paragraphs[:25])[:1500]
    return {"meta": meta, "headings": heads, "text": body}

# ---------- IMAGE ----------
from PIL import Image, ExifTags
EXIF_T = {v: k for k, v in ExifTags.TAGS.items()}

def image_facts(p):
    im = Image.open(p)
    w, h = im.size
    ex = {}
    try:
        raw = im.getexif()
        for tag in ("Make", "Model", "DateTimeOriginal", "DateTime", "LensModel"):
            tid = EXIF_T.get(tag)
            if tid and raw.get(tid): ex[tag] = str(raw.get(tid))[:60]
        if raw.get_ifd(0x8825): ex["GPS"] = True
    except Exception: pass
    # screenshot heuristic: no camera, PNG, screen-ish aspect
    is_shot = (not ex.get("Make")) and p.suffix.lower() in (".png", ".jpeg", ".jpg") and \
              (max(w, h) / min(w, h) < 2.4) and max(w, h) >= 640
    return {"size": [w, h], "exif": ex, "camera": bool(ex.get("Make")), "maybe_screenshot": is_shot}


HANDLERS = {"pdf": pdf_facts, "docx": docx_facts,
            "png": image_facts, "jpg": image_facts, "jpeg": image_facts,
            "heic": image_facts, "webp": image_facts, "tif": image_facts}

stats, samples, fails = {}, {}, Counter()
LIMIT = int(sys.argv[1]) if len(sys.argv) > 1 else 60

for ext, fn in HANDLERS.items():
    sel = [p for p in files if p.suffix.lower().lstrip(".") == ext][:LIMIT]
    if not sel: continue
    ok, t0 = 0, time.time()
    got = []
    for p in sel:
        try:
            r = fn(p); ok += 1
            if len(got) < 3: got.append((p.name, r))
        except Exception as e:
            fails[f"{ext}:{type(e).__name__}"] += 1
    dt = time.time() - t0
    stats[ext] = {"n": len(sel), "ok": ok, "total_s": round(dt, 2),
                  "ms_per_file": round(1000 * dt / max(len(sel), 1), 1)}
    samples[ext] = got
    print(f"{ext:<6} n={len(sel):<4} ok={ok:<4} {1000*dt/max(len(sel),1):7.1f} ms/file  total {dt:5.2f}s", flush=True)

print("\n=== failures ===", dict(fails), flush=True)
print("\n=== WHAT WE ACTUALLY LEARN (samples) ===\n", flush=True)
for ext, got in samples.items():
    for name, r in got[:2]:
        print(f"--- {name[:70]}  [{ext}]")
        if ext in ("pdf", "docx"):
            if r.get("meta"): print(f"    meta: {r['meta']}")
            if r.get("headings"): print(f"    headings: {r['headings'][:4]}")
            t = re.sub(r"\s+", " ", r.get("text", ""))[:260]
            print(f"    text: {t}")
        else:
            print(f"    {r}")
        print(flush=True)

json.dump(stats, open("extract_stats.json", "w"), indent=2)
