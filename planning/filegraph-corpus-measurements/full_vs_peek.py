"""Peek (page 1 + metadata) vs FULL document text: what does it cost, how much more do we learn?"""
import time, re, statistics
from pathlib import Path
from collections import Counter

DL = Path.home() / "Downloads"
import pypdfium2 as pdfium
from docx import Document

pdfs = [p for p in DL.iterdir() if p.is_file() and p.suffix.lower() == ".pdf"][:120]
docxs = [p for p in DL.iterdir() if p.is_file() and p.suffix.lower() == ".docx"][:80]
print(f"PDFs: {len(pdfs)}   DOCXs: {len(docxs)}\n", flush=True)


def pdf_pages(p, limit=None):
    d = pdfium.PdfDocument(str(p))
    n = len(d)
    take = n if limit is None else min(limit, n)
    t = []
    for i in range(take):
        t.append(d[i].get_textpage().get_text_range())
    d.close()
    return "".join(t), n


# ---------- PDF: peek vs full ----------
peek_t = full_t = 0.0
peek_chars = full_chars = 0
pages_all, empty_peek, empty_full, rescued = [], 0, 0, 0

for p in pdfs:
    try:
        t0 = time.time(); tx, n = pdf_pages(p, 1); peek_t += time.time() - t0
        peek_chars += len(tx.strip()); pe = len(tx.strip()) < 50
        t0 = time.time(); tf, _ = pdf_pages(p, None); full_t += time.time() - t0
        full_chars += len(tf.strip()); fe = len(tf.strip()) < 50
        pages_all.append(n)
        empty_peek += pe; empty_full += fe
        if pe and not fe: rescued += 1
    except Exception:
        pass

n = len(pages_all)
print("=== PDF ===")
print(f"pages: median {statistics.median(pages_all):.0f}  mean {statistics.mean(pages_all):.1f}  max {max(pages_all)}")
print(f"PEEK (page 1) : {1000*peek_t/n:6.1f} ms/file   {peek_chars//n:7,} chars/file avg")
print(f"FULL (all pgs): {1000*full_t/n:6.1f} ms/file   {full_chars//n:7,} chars/file avg")
print(f"  cost  x{full_t/max(peek_t,1e-9):.1f}      text  x{full_chars/max(peek_chars,1):.1f}")
print(f"  files with NO usable text: peek {empty_peek}/{n} ({100*empty_peek/n:.0f}%)  full {empty_full}/{n} ({100*empty_full/n:.0f}%)")
print(f"  RESCUED by reading past page 1: {rescued} files\n", flush=True)

# ---------- DOCX: peek vs full ----------
pt = ft = 0.0; pc = fc = 0; dn = 0
for p in docxs:
    try:
        t0 = time.time(); d = Document(str(p)); head = " ".join(x.text for x in d.paragraphs[:25]); pt += time.time() - t0
        t0 = time.time(); d2 = Document(str(p))
        body = " ".join(x.text for x in d2.paragraphs)
        for tb in d2.tables:
            for row in tb.rows:
                body += " " + " ".join(c.text for c in row.cells)
        ft += time.time() - t0
        pc += len(head.strip()); fc += len(body.strip()); dn += 1
    except Exception:
        pass
print("=== DOCX ===")
print(f"PEEK (25 paras): {1000*pt/max(dn,1):6.1f} ms/file  {pc//max(dn,1):7,} chars/file")
print(f"FULL (+tables) : {1000*ft/max(dn,1):6.1f} ms/file  {fc//max(dn,1):7,} chars/file")
print(f"  cost x{ft/max(pt,1e-9):.1f}   text x{fc/max(pc,1):.1f}\n", flush=True)

# ---------- projected corpus cost ----------
N_PDF, N_DOCX = 811, 331
print("=== projected on your full corpus ===")
print(f"PEEK: {(N_PDF*peek_t/n + N_DOCX*pt/max(dn,1)):6.1f} s")
print(f"FULL: {(N_PDF*full_t/n + N_DOCX*ft/max(dn,1)):6.1f} s", flush=True)
