"""HOLD-OUT TEST: can the graph recover a fact the file itself doesn't reveal?
For every file where rules confidently found `school`, hide it, then propagate from neighbours.
Measures precision/recall of propagation — the least-measured mechanism in the spec."""
import re, time, math
from pathlib import Path
from collections import defaultdict, Counter

DL = Path.home() / "Downloads"
import pypdfium2 as pdfium
from docx import Document

SCHOOLS = {
    "columbia": "Columbia", "barnard": "Barnard", "yale": "Yale", "stanford": "Stanford",
    "harvard": "Harvard", "northwestern": "Northwestern", "duke": "Duke", "dartmouth": "Dartmouth",
    "uchicago": "UChicago", "u chicago": "UChicago", "university of chicago": "UChicago",
    "cornell": "Cornell", "johns hopkins": "Johns Hopkins", "georgetown": "Georgetown",
    "wash u": "WashU", "washu": "WashU", "washington university": "WashU", "hku": "HKU",
    "unc": "UNC", "princeton": "Princeton", "mit": "MIT", "nyu": "NYU", "upenn": "UPenn",
}
def wb(k): return re.compile(r'(?<![a-z0-9])' + re.escape(k) + r'(?![a-z0-9])', re.I)
SRE = {k: (wb(k), v) for k, v in SCHOOLS.items()}

ENT = {
    "email": re.compile(r'\b[\w.+-]+@([\w-]+\.[\w.-]+)\b'),
    "code":  re.compile(r'\b([A-Z]{2,5})\s?-?\s?(\d{3,4})\b'),
    "uniid": re.compile(r'\b([a-z]{2,4}\d{4})\b'),
    "doi":   re.compile(r'\b10\.\d{4,9}/[-._;()/:A-Za-z0-9]+'),
}
JUNK_DOM = {"gmail.com", "googlemail.com", "outlook.com", "hotmail.com", "yahoo.com",
            "icloud.com", "w3.org", "adobe.com", "openxmlformats.org"}


def read(p):
    try:
        if p.suffix.lower() == ".pdf":
            d = pdfium.PdfDocument(str(p)); t = [d[i].get_textpage().get_text_range() for i in range(len(d))]
            d.close(); return "".join(t)
        doc = Document(str(p))
        b = " ".join(x.text for x in doc.paragraphs)
        for tb in doc.tables:
            for r in tb.rows: b += " " + " ".join(c.text for c in r.cells)
        return b
    except Exception: return None


def school_of(name, text):
    """disciplined extractor: word boundary + position weight + rank + threshold"""
    sc = defaultdict(float)
    for zt, w in ((name, 10.0), (text[:1200], 3.0), (text[1200:], 0.4)):
        for k, (rx, v) in SRE.items():
            n = len(rx.findall(zt))
            if n: sc[v] += w * (1 + 0.15 * min(n, 6))
    if not sc: return None
    r = sorted(sc.items(), key=lambda kv: -kv[1])
    top = r[0][1]; second = r[1][1] if len(r) > 1 else 0.0
    return r[0][0] if (top >= 3.0 and top - second >= 1.0) else None


def entities(name, text):
    out = set(); hay = f"{name}\n{text[:8000]}"
    for m in ENT["email"].finditer(hay):
        d = m.group(1).lower()
        if d not in JUNK_DOM: out.add(("email", d))
    for m in ENT["code"].finditer(hay):  out.add(("code", f"{m.group(1)}{m.group(2)}"))
    for m in ENT["uniid"].finditer(hay): out.add(("uniid", m.group(1)))
    for m in ENT["doi"].finditer(hay):   out.add(("doi", m.group(0)[:40]))
    return out


files = [p for p in DL.iterdir() if p.is_file() and p.suffix.lower() in (".pdf", ".docx")][:400]
print(f"reading {len(files)} files…", flush=True)
t0 = time.time()
truth, ents, mt, names = {}, {}, {}, {}
for i, p in enumerate(files):
    t = read(p)
    if t is None: continue
    names[i] = p.name; mt[i] = p.stat().st_mtime
    truth[i] = school_of(p.name, t)
    ents[i] = entities(p.name, t)
idx = list(names)
print(f"read in {time.time()-t0:.0f}s   known-school files: {sum(1 for i in idx if truth[i])}/{len(idx)}\n", flush=True)

# ---- build fact edges with IDF weighting + hub suppression ----
df = Counter()
for i in idx: df.update(ents[i])
N = len(idx)
KEEP = {e for e, c in df.items() if 2 <= c <= max(3, int(0.25 * N))}   # hub cap
inv = defaultdict(list)
for i in idx:
    for e in ents[i] & KEEP: inv[e].append(i)

W = defaultdict(float)
for e, fl in inv.items():
    w = math.log(N / len(fl))                       # IDF
    for a in range(len(fl)):
        for b in range(a + 1, len(fl)):
            W[(fl[a], fl[b])] += w
# session co-occurrence: same download window
srt = sorted(idx, key=lambda i: mt[i])
for a in range(len(srt) - 1):
    for b in range(a + 1, min(a + 6, len(srt))):
        if abs(mt[srt[a]] - mt[srt[b]]) <= 600:      # 10 minutes
            W[(min(srt[a], srt[b]), max(srt[a], srt[b]))] += 0.35

adj = defaultdict(list)
for (a, b), w in W.items():
    adj[a].append((b, w)); adj[b].append((a, w))
print(f"edges: {len(W)}  avg degree {2*len(W)/N:.1f}\n", flush=True)

# ---- HOLD-OUT: hide each known school, recover from neighbours ----
for MIN_W in (0.5, 1.0, 2.0):
    ok = wrong = abst = 0
    examples = []
    for i in idx:
        if not truth[i]: continue
        vote = defaultdict(float)
        for j, w in adj[i]:
            if j != i and truth[j] and w >= MIN_W:   # neighbours keep their labels
                vote[truth[j]] += w
        if not vote: abst += 1; continue
        r = sorted(vote.items(), key=lambda kv: -kv[1])
        pred, top = r[0]
        second = r[1][1] if len(r) > 1 else 0.0
        if top < MIN_W or (top - second) < 0.5 * top:    # margin gate
            abst += 1; continue
        if pred == truth[i]:
            ok += 1
        else:
            wrong += 1
            if len(examples) < 6: examples.append((names[i], truth[i], pred, round(top, 1)))
    tot = ok + wrong
    prec = 100 * ok / tot if tot else 0
    print(f"MIN_W={MIN_W}:  correct {ok}  wrong {wrong}  abstained {abst}   "
          f"precision {prec:.0f}%   coverage {100*tot/max(1,sum(1 for i in idx if truth[i])):.0f}%")
    for n, t_, p_, w_ in examples:
        print(f"      WRONG  {n[:52]:<54} true={t_:<14} got={p_} (w={w_})")
    print(flush=True)
