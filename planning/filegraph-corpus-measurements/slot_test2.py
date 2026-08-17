"""v2: word boundaries + positional scoring + rank-don't-first-match + per-slot abstention.
How much of the v1 error was sloppiness rather than a missing neural model?"""
import re, time
from pathlib import Path
from collections import Counter, defaultdict

DL = Path.home() / "Downloads"
import pypdfium2 as pdfium
from docx import Document

SCHOOLS = {
    "columbia": "Columbia", "barnard": "Barnard", "yale": "Yale", "stanford": "Stanford",
    "harvard": "Harvard", "northwestern": "Northwestern", "duke": "Duke", "dartmouth": "Dartmouth",
    "uchicago": "UChicago", "u chicago": "UChicago", "university of chicago": "UChicago",
    "cornell": "Cornell", "johns hopkins": "Johns Hopkins", "georgetown": "Georgetown",
    "wash u": "WashU", "washu": "WashU", "washington university": "WashU",
    "hku": "HKU", "university of hong kong": "HKU", "unc": "UNC", "princeton": "Princeton",
    "mit": "MIT", "nyu": "NYU", "upenn": "UPenn", "morehead-cain": "UNC",
}
WORKTYPE = {
    "syllabus": "Syllabus", "lecture": "Lecture", "problem set": "Problem Set", "pset": "Problem Set",
    "homework": "Homework", "assignment": "Assignment", "midterm": "Exam", "final exam": "Exam",
    "exam": "Exam", "quiz": "Quiz", "solution": "Solutions", "essay": "Essay",
    "lab report": "Lab Report", "transcript": "Transcript", "supplement": "Application",
    "application": "Application", "resume": "Resume", "cover letter": "Cover Letter",
    "abstract": "Abstract", "manuscript": "Manuscript", "invoice": "Invoice", "receipt": "Receipt",
}
SUBJ = re.compile(r'\b([A-Z]{2,5})\s?-?\s?(\d{3,4})(?:\.\d{1,4})?\b')
ACADEMIC_CTX = re.compile(r'\b(course|section|syllabus|lecture|semester|instructor|professor|'
                          r'credits|prerequisite|department|registrar|enroll|class number)\b', re.I)
MONTHS = re.compile(r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b', re.I)
TERM = re.compile(r'\b(Spring|Fall|Autumn|Summer|Winter)\s*(20\d{2})\b', re.I)
YEAR = re.compile(r'\b(20[0-2]\d)\b')

# v2 FIX 1: word-boundary matcher, built once per gazetteer key
def wb(k):
    return re.compile(r'(?<![a-z0-9])' + re.escape(k) + r'(?![a-z0-9])', re.I)
SCHOOL_RE = {k: (wb(k), v) for k, v in SCHOOLS.items()}
WORK_RE = {k: (wb(k), v) for k, v in WORKTYPE.items()}

MIN_SCORE = 3.0   # v2 FIX 4: abstain below this


def read(p):
    try:
        if p.suffix.lower() == ".pdf":
            d = pdfium.PdfDocument(str(p)); t = [d[i].get_textpage().get_text_range() for i in range(len(d))]
            d.close(); return "".join(t)
        doc = Document(str(p))
        b = " ".join(x.text for x in doc.paragraphs)
        for tb in doc.tables:
            for r in tb.rows: b += " " + " ".join(c.text for c in r.cells)
        return b
    except Exception:
        return None


def zones(name, text):
    """v2 FIX 3: position matters. filename >> head >> body."""
    return [(name, 10.0), (text[:1200], 3.0), (text[1200:], 0.4)]


def best(cands):
    """v2 FIX 2: rank, don't take the first. Returns (value, score, margin)."""
    if not cands: return None, 0.0, 0.0
    r = sorted(cands.items(), key=lambda kv: -kv[1])
    top = r[0][1]; second = r[1][1] if len(r) > 1 else 0.0
    return r[0][0], top, top - second


def slots(name, text):
    s, ev = {}, {}
    Z = zones(name, text)

    sc = defaultdict(float)
    for zt, w in Z:
        for k, (rx, v) in SCHOOL_RE.items():
            n = len(rx.findall(zt))
            if n: sc[v] += w * (1 + 0.15 * min(n, 6))
    v, top, marg = best(sc)
    if v and top >= MIN_SCORE and marg >= 1.0:
        s["school"] = v; ev["school"] = f"score={top:.1f} margin={marg:.1f}"

    hay = f"{name}\n{text[:4000]}"
    m = TERM.search(hay)
    if m:
        s["term"] = f"{m.group(2)}-{m.group(1).title()}"; ev["term"] = f"explicit:{m.group(0)}"
    else:
        yrs = YEAR.findall(name) or YEAR.findall(text[:1500])
        if yrs:
            c = Counter(yrs).most_common(1)[0]
            if c[1] >= 2 or yrs is YEAR.findall(name):
                s["term"] = c[0]; ev["term"] = f"year:{c[0]}"

    for m in SUBJ.finditer(f"{name}\n{text[:6000]}"):
        # v2 FIX 5: reject month-like prefixes (JUN 202), require academic context
        if MONTHS.fullmatch(m.group(1)): continue
        a, b = max(0, m.start() - 250), m.end() + 250
        if ACADEMIC_CTX.search(f"{name}\n{text[:6000]}"[a:b]):
            s["subject"] = f"{m.group(1)}{m.group(2)}"; ev["subject"] = f"code+ctx:{m.group(0)}"
            break

    wc = defaultdict(float)
    for zt, w in [(name, 10.0), (text[:800], 2.0)]:   # v2 FIX 6: never the whole body
        for k, (rx, v) in WORK_RE.items():
            if rx.search(zt): wc[v] += w
    v, top, marg = best(wc)
    if v and top >= MIN_SCORE:
        s["work_type"] = v; ev["work_type"] = f"score={top:.1f}"
    return s, ev


files = [p for p in DL.iterdir() if p.is_file() and p.suffix.lower() in (".pdf", ".docx")][:300]
t0 = time.time(); filled = Counter(); rows = []; n = 0; got = {}
for p in files:
    t = read(p)
    if t is None: continue
    n += 1
    s, ev = slots(p.name, t)
    got[p.name] = s
    for k in s: filled[k] += 1
    if len(s) >= 2: rows.append((p.name, s, ev))
dt = time.time() - t0

print(f"v2: {n} files in {dt:.1f}s ({1000*dt/n:.0f} ms/file)\n")
print("=== FILL RATE  (v1 → v2) ===")
for k, v1 in (("school", 75), ("term", 70), ("subject", 12), ("work_type", 70)):
    print(f"  {k:<10} {100*filled[k]/n:3.0f}%   (v1 {v1}%)")
print()
print("=== the files v1 got WRONG ===")
for probe in ["Wash U .docx", "U Chicago Supplemental essay 1 (2).pdf",
              "Probability For Engineers.pdf", "Adobe Scan Jun 12, 2025 (2).pdf",
              "Stanford (2).docx", "Syllabus BUSIB 4300 Spring 2026 Haran Segram.pdf",
              "HKU Continued Interest.docx", "Duke Supplemental.docx"]:
    if probe in got:
        s = got[probe]
        path = "/".join(s.get(k, "·") for k in ("school", "term", "subject", "work_type"))
        print(f"  {probe[:52]:<54} → {path}")
