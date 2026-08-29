# src/readers/docx_python_docx.py
"""`read_docx` backed by python-docx.

**Why this exists at all.** `deployment.py` wired `read_docx = _no_reader`, whose
docstring says *"this deployment ships no library for the format"*. That sentence
was true when it was written and is not true now: `python-docx` is installed in
this interpreter. The consequence of leaving it was not a missing feature but a
LIE IN THE RECORD -- every `.docx` on a person's disk recorded `unsupported`,
which §2.4 defines as *"no reader exists and the bytes were never looked at"*, and
downstream every count agreed that those files carried nothing. On a real human's
disk that is most of the writing they have ever done.

**What is library knowledge and what is not.** `Region`'s contract: *"the reader
says WHAT KIND OF PLACE this is, because that is library knowledge (a heading
style, a table cell, a footer)"*. python-docx exposes the paragraph STYLE, which is
what Word itself uses to mean "this is a heading", and the outline LEVEL that goes
with it. So the zone here is read from the style, never guessed from the text --
no "short line in title case is probably a heading", which would be exactly the
structural judgement P4 forbids a reader to invent.

**What this reader does not claim.** No links, relationships, annotations or
revision metadata. `DocxDocument` has slots for all of them and they stay empty,
because an empty tuple from a reader that never looked and an empty tuple from a
document that has none are the same value and this module will not pretend
otherwise by filling them badly. They are the honest next increment, not a gap
hidden behind a default.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any, Callable

import docx
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from extractors.docx import DocxCell, DocxDocument, DocxParagraph

#: The core properties P4 keeps as strings. `title` is its own zone (P5
#: `docx.py`'s `TITLE_PROPERTIES`); the rest travel as document metadata. Read by
#: NAME rather than by iterating the object, because python-docx exposes a fixed
#: set of attributes and a `dir()` walk would pick up methods.
_CORE_PROPERTIES: tuple[str, ...] = (
    "title", "author", "subject", "keywords", "category", "comments",
    "last_modified_by", "content_status", "identifier", "language", "version",
)

#: The core properties whose value is a datetime rather than a string. §2.3 keeps
#: them apart from the text slots, and P4 D8 wants them ISO-8601.
_DATE_PROPERTIES: tuple[str, ...] = ("created", "modified", "last_printed")

#: Word's own names for the built-in heading styles, and the only thing consulted
#: to decide that a paragraph is a heading. A document that renames them loses the
#: zone rather than gaining a guessed one.
_HEADING_STYLE = "heading"
_TITLE_STYLE = "title"


def _zone(paragraph: Paragraph) -> tuple[str, int | None]:
    """P4's zone for this paragraph, and its outline depth if it is a heading.

    `style.name` rather than the raw style id: python-docx resolves the id through
    the document's style table, so a document that inherits `Heading 1` from a
    template still reads as a heading. A paragraph inside a header or footer part
    never reaches this function -- python-docx exposes those on a different object
    and this reader does not walk them, which is why `header_footer` is not
    produced here rather than being produced wrongly.
    """
    name = (paragraph.style.name or "").strip().lower() if paragraph.style else ""
    if name.startswith(_HEADING_STYLE):
        suffix = name[len(_HEADING_STYLE):].strip()
        # `Heading 1` -> 1. An unnumbered `Heading` is still a heading; it sits at
        # the outermost level, which is what Word renders it as.
        return "heading", int(suffix) if suffix.isdigit() else 1
    if name == _TITLE_STYLE:
        return "heading", 1
    return "body", None


def _body_blocks(document: _Document):
    """Paragraphs and tables in the order the document lays them out.

    python-docx's `document.paragraphs` skips everything inside a table and
    `document.tables` loses where each table sat, so neither alone can say what
    came before what. The body element's own child order is the only place that
    ordering exists, and P4 anchors spans to a paragraph ordinal -- an ordinal
    assigned in the wrong order would point every citation at the wrong text.
    """
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _read(path: Path) -> DocxDocument | None:
    document = docx.Document(str(path))

    core = document.core_properties
    properties: dict[str, str] = {}
    for name in _CORE_PROPERTIES:
        value = getattr(core, name, None)
        if value:
            properties[name] = str(value)
    iso_dates: dict[str, str] = {}
    for name in _DATE_PROPERTIES:
        value = getattr(core, name, None)
        if value is not None:
            iso_dates[name] = value.isoformat()

    paragraphs: list[DocxParagraph] = []
    cells: list[DocxCell] = []
    #: The heading ancestry as (ordinal, label) pairs, outermost first. A heading
    #: at level N replaces everything from N down, which is what nesting means.
    ancestry: list[tuple[int, int, str]] = []  # (level, ordinal, label)
    #: ONE-BASED, because P4 D3 says container-path indices are: a `paragraph`
    #: segment at index 0 raises `container-path indices are 1-based`, and the
    #: whole extraction is recorded `failed` -- a real document reported as a
    #: damaged one. Every ordinal below is derived from this counter, so the
    #: heading ancestry is 1-based with it and no second convention exists.
    index = 1
    table_index = 1

    for block in _body_blocks(document):
        if isinstance(block, Table):
            #: Row 0 is the header row. That is a convention rather than a fact
            #: Word records -- `w:tblHeader` marks a REPEATING header and most
            #: documents that have a header row do not set it -- so it is applied
            #: only to name the column, never to drop the row: row 0 is emitted as
            #: a cell like any other and a caller that disagrees still has it.
            rows = block.rows
            headers: list[str | None] = []
            if rows:
                headers = [cell.text.strip() or None for cell in rows[0].cells]
            for row_index, row in enumerate(rows, start=1):
                for column_index, cell in enumerate(row.cells, start=1):
                    header = (headers[column_index - 1]
                              if column_index <= len(headers) else None)
                    cells.append(DocxCell(
                        table=table_index, row=row_index, column=column_index,
                        text=cell.text,
                        # A header row is not its own column header.
                        column_header=None if row_index == 1 else header))
            table_index += 1
            continue

        text = block.text
        if not text.strip():
            # An empty paragraph is layout, not content. It still consumes no
            # ordinal, so the ordinals stay dense and a person counting
            # paragraphs in Word and a citation here agree about which is which.
            continue
        zone, level = _zone(block)
        if zone == "heading" and level is not None:
            ancestry = [entry for entry in ancestry if entry[0] < level]
            ancestry.append((level, index, text))
        paragraphs.append(DocxParagraph(
            index=index, text=text, zone=zone,
            heading_path=tuple((ordinal, label)
                               for _, ordinal, label in ancestry)))
        index += 1

    return DocxDocument(
        core_properties=properties, paragraphs=tuple(paragraphs),
        cells=tuple(cells), iso_dates=iso_dates)


def python_docx_reader() -> Callable[[Path], DocxDocument | None]:
    """The wired `read_docx`.

    A factory for the same reason `pdfminer_reader` is one: it is the seam a
    deployment swaps, and a bare function would make the call site look like it
    had chosen a library rather than been given one.
    """

    def read_docx(path: Path, **_: Any) -> DocxDocument | None:
        return _read(Path(path))

    return read_docx
